"""
文档解析器 - 解析各种格式的文档

支持格式：
- PDF
- Word (.docx)
- Excel (.xlsx)
- 文本文件 (.txt, .md)
"""

import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from abc import ABC, abstractmethod

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """文档解析器抽象基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析结果字典
        """
        pass


class PDFParser(DocumentParser):
    """PDF 文档解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析 PDF 文档"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            
            # 提取文本
            text_content = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_content.append({
                        "page": page_num,
                        "content": page_text.strip()
                    })
            
            # 提取元数据
            metadata = reader.metadata if reader.metadata else {}
            
            return {
                "success": True,
                "type": "pdf",
                "pages": len(reader.pages),
                "content": text_content,
                "full_text": "\n\n".join([page["content"] for page in text_content]),
                "metadata": {
                    "title": metadata.get("/Title", ""),
                    "author": metadata.get("/Author", ""),
                    "subject": metadata.get("/Subject", ""),
                    "creator": metadata.get("/Creator", ""),
                }
            }
            
        except ImportError:
            logger.error("❌ PyPDF2 未安装，请运行: pip install PyPDF2")
            return {"success": False, "error": "PyPDF2 未安装"}
        except Exception as e:
            logger.error(f"❌ PDF 解析失败: {e}")
            return {"success": False, "error": str(e)}


class WordParser(DocumentParser):
    """Word 文档解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析 Word 文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # 提取核心属性
            core_props = doc.core_properties
            
            return {
                "success": True,
                "type": "word",
                "paragraphs": paragraphs,
                "tables": tables,
                "full_text": "\n\n".join(paragraphs),
                "metadata": {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else ""
                }
            }
            
        except ImportError:
            logger.error("❌ python-docx 未安装，请运行: pip install python-docx")
            return {"success": False, "error": "python-docx 未安装"}
        except Exception as e:
            logger.error(f"❌ Word 解析失败: {e}")
            return {"success": False, "error": str(e)}


class ExcelParser(DocumentParser):
    """Excel 文档解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析 Excel 文档"""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(file_path, data_only=True)
            
            # 提取所有工作表
            sheets = {}
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                # 提取数据
                data = []
                for row in sheet.iter_rows(values_only=True):
                    # 跳过完全空的行
                    if any(cell is not None for cell in row):
                        data.append(list(row))
                
                sheets[sheet_name] = data
            
            # 生成摘要文本
            summary_parts = []
            for sheet_name, data in sheets.items():
                if data:
                    summary_parts.append(f"工作表: {sheet_name}")
                    summary_parts.append(f"行数: {len(data)}")
                    if data:
                        summary_parts.append(f"列数: {len(data[0])}")
                    summary_parts.append("")
            
            return {
                "success": True,
                "type": "excel",
                "sheets": sheets,
                "sheet_names": wb.sheetnames,
                "summary": "\n".join(summary_parts),
                "metadata": {
                    "sheet_count": len(wb.sheetnames)
                }
            }
            
        except ImportError:
            logger.error("❌ openpyxl 未安装，请运行: pip install openpyxl")
            return {"success": False, "error": "openpyxl 未安装"}
        except Exception as e:
            logger.error(f"❌ Excel 解析失败: {e}")
            return {"success": False, "error": str(e)}


class TextParser(DocumentParser):
    """文本文件解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析文本文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            if content is None:
                raise ValueError("无法使用任何编码读取文件")
            
            # 统计行数
            lines = content.splitlines()
            non_empty_lines = [line for line in lines if line.strip()]
            
            return {
                "success": True,
                "type": "text",
                "content": content,
                "lines": len(lines),
                "non_empty_lines": len(non_empty_lines),
                "encoding": used_encoding,
                "metadata": {}
            }
            
        except Exception as e:
            logger.error(f"❌ 文本解析失败: {e}")
            return {"success": False, "error": str(e)}


class DocumentParserFactory:
    """文档解析器工厂"""
    
    # 文件扩展名到解析器的映射
    PARSERS = {
        ".pdf": PDFParser,
        ".docx": WordParser,
        ".doc": WordParser,
        ".xlsx": ExcelParser,
        ".xls": ExcelParser,
        ".txt": TextParser,
        ".md": TextParser,
        ".markdown": TextParser,
    }
    
    @classmethod
    def get_parser(cls, file_path: str) -> Optional[DocumentParser]:
        """
        根据文件类型获取解析器
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析器实例，如果不支持则返回 None
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        parser_class = cls.PARSERS.get(extension)
        if parser_class:
            return parser_class()
        
        logger.warning(f"⚠️  不支持的文件类型: {extension}")
        return None
    
    @classmethod
    def parse_file(cls, file_path: str) -> Dict[str, Any]:
        """
        解析文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析结果
        """
        parser = cls.get_parser(file_path)
        if not parser:
            return {
                "success": False,
                "error": f"不支持的文件类型: {Path(file_path).suffix}"
            }
        
        return parser.parse(file_path)
    
    @classmethod
    def supported_extensions(cls) -> List[str]:
        """获取支持的文件扩展名列表"""
        return list(cls.PARSERS.keys())


# 便捷函数
def parse_document(file_path: str) -> Dict[str, Any]:
    """
    解析文档
    
    Args:
        file_path: 文件路径
        
    Returns:
        解析结果
    """
    return DocumentParserFactory.parse_file(file_path)

